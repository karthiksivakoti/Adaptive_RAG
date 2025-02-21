# risk_rag_system/input_processing/document/parsers/docx_parser.py

from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
import io
from PIL import Image
import pytesseract
from loguru import logger
from pydantic import BaseModel
from datetime import datetime

class DOCXParserConfig(BaseModel):
    """Configuration for DOCX parsing"""
    extract_images: bool = True
    perform_ocr: bool = True
    preserve_formatting: bool = True
    extract_comments: bool = True
    extract_headers_footers: bool = True
    extract_tables: bool = True
    image_quality: int = 95
    ocr_lang: str = "eng"

class DOCXParser:
    """Parser for DOCX documents"""
    
    def __init__(self, config: Optional[DOCXParserConfig] = None):
        self.config = config or DOCXParserConfig()
        logger.info("Initialized DOCXParser")

    async def parse(
        self,
        file_path: Path,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            logger.info(f"Parsing DOCX: {file_path}")
            doc = docx.Document(file_path)
            
            # Extract content
            content = []
            images = []
            tables = []
            
            # Process headers if enabled
            if self.config.extract_headers_footers:
                header_text = self._extract_headers(doc)
                if header_text:
                    content.append(header_text)
            
            # Process main content
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Process paragraph
                    paragraph = Paragraph(element, doc)
                    text = self._process_paragraph(paragraph)
                    if text:
                        content.append(text)
                
                elif isinstance(element, CT_Tbl):
                    # Process table
                    table = Table(element, doc)
                    if self.config.extract_tables:
                        table_data = self._process_table(table)
                        if table_data:
                            tables.append(table_data)
            
            # Process footers if enabled
            if self.config.extract_headers_footers:
                footer_text = self._extract_footers(doc)
                if footer_text:
                    content.append(footer_text)
            
            # Extract images if enabled
            if self.config.extract_images:
                images = await self._extract_images(doc)
            
            # Process images with OCR if enabled
            if self.config.perform_ocr and images:
                ocr_text = await self._process_ocr(images)
                if ocr_text:
                    content.append(ocr_text)
            
            # Get document properties
            doc_props = self._get_document_properties(doc)
            
            # Prepare result
            result = {
                "content": "\n\n".join(content),
                "metadata": {
                    "source": str(file_path),
                    "type": "docx",
                    **doc_props,
                    **(metadata or {})
                },
                "images": images,
                "tables": tables
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise

    def _process_paragraph(self, paragraph: Paragraph) -> str:
        """Process paragraph with formatting"""
        if not paragraph.text.strip():
            return ""

        if not self.config.preserve_formatting:
            return paragraph.text

        formatted_text = []
        for run in paragraph.runs:
            text = run.text
            if text.strip():
                # Apply formatting
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                if run.underline:
                    text = f"__{text}__"
                formatted_text.append(text)

        # Handle paragraph styles
        result = "".join(formatted_text)
        style_name = paragraph.style.name.lower()
        
        if "heading" in style_name:
            level = int(style_name[-1]) if style_name[-1].isdigit() else 1
            result = f"{'#' * level} {result}"
        elif "list" in style_name:
            result = f"- {result}"

        return result

    def _process_table(self, table: Table) -> Dict[str, Any]:
        """Process table from document"""
        try:
            data = []
            rows = list(table.rows)
            
            for row in rows:
                row_data = []
                for cell in row.cells:
                    # Combine all paragraphs in cell
                    cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    row_data.append(cell_text)
                data.append(row_data)

            return {
                "data": data,
                "num_rows": len(data),
                "num_cols": len(data[0]) if data else 0,
                "style": "table"
            }
            
        except Exception as e:
            logger.error(f"Error processing table: {e}")
            return None

    async def _extract_images(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract images from document"""
        try:
            images = []
            rels = doc.part.rels
            
            for rel in rels.values():
                if "image" in rel.reltype:
                    try:
                        image_bytes = rel.target_part.blob
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        images.append({
                            "content": image_bytes,
                            "size": image.size,
                            "format": image.format.lower(),
                            "index": len(images)
                        })
                        
                    except Exception as img_error:
                        logger.warning(f"Error processing image: {img_error}")
                        continue

            return images
            
        except Exception as e:
            logger.error(f"Error extracting images: {e}")
            return []

    async def _process_ocr(self, images: List[Dict[str, Any]]) -> str:
        """Process OCR on extracted images"""
        try:
            ocr_results = []
            
            for img_data in images:
                image = Image.open(io.BytesIO(img_data["content"]))
                text = pytesseract.image_to_string(
                    image,
                    lang=self.config.ocr_lang
                )
                
                if text.strip():
                    ocr_results.append(text)

            return "\n\n".join(ocr_results)
            
        except Exception as e:
            logger.error(f"Error processing OCR: {e}")
            return ""

    def _extract_headers(self, doc: Document) -> str:
        """Extract text from headers"""
        try:
            header_texts = []
            for section in doc.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            header_texts.append(self._process_paragraph(paragraph))
            
            return "\n".join(header_texts)
            
        except Exception as e:
            logger.error(f"Error extracting headers: {e}")
            return ""

    def _extract_footers(self, doc: Document) -> str:
        """Extract text from footers"""
        try:
            footer_texts = []
            for section in doc.sections:
                footer = section.footer
                if footer:
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            footer_texts.append(self._process_paragraph(paragraph))
            
            return "\n".join(footer_texts)
            
        except Exception as e:
            logger.error(f"Error extracting footers: {e}")
            return ""

    def _get_document_properties(self, doc: Document) -> Dict[str, Any]:
        """Extract document properties"""
        try:
            core_properties = doc.core_properties
            return {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "created": core_properties.created.isoformat() if core_properties.created else "",
                "modified": core_properties.modified.isoformat() if core_properties.modified else "",
                "last_modified_by": core_properties.last_modified_by or "",
                "revision": core_properties.revision or 1
            }
        except Exception as e:
            logger.error(f"Error getting document properties: {e}")
            return {}
                