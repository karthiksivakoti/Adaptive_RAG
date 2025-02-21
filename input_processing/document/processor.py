# risk_rag_system/input_processing/document/processor.py

from typing import Dict, Any, List, Optional, BinaryIO
from pathlib import Path
import fitz  # PyMuPDF
import docx
import pandas as pd
import time
import pytesseract
from PIL import Image
import io
from loguru import logger
from pydantic import BaseModel
import numpy as np
from concurrent.futures import ThreadPoolExecutor
import hashlib

class ProcessedDocument(BaseModel):
    """Structure for processed document content"""
    content: str
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    file_hash: str

class ProcessorConfig(BaseModel):
    """Configuration for document processor"""
    supported_formats: List[str] = [
        ".pdf", ".docx", ".doc", 
        ".xlsx", ".xls", 
        ".txt", ".jpg", ".png"
    ]
    extract_tables: bool = True
    ocr_enabled: bool = True
    preserve_formatting: bool = True
    max_file_size_mb: int = 50
    tesseract_path: Optional[str] = None
    max_workers: int = 4

class DocumentProcessor:
    """Handles document processing and content extraction"""
    
    def __init__(self, config: Optional[ProcessorConfig] = None):
        self.config = config or ProcessorConfig()
        if self.config.tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = self.config.tesseract_path
        self.executor = ThreadPoolExecutor(max_workers=self.config.max_workers)
        logger.info("Initialized DocumentProcessor")

    async def process_file(
        self,
        file_path: Path,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process a file and extract content"""
        try:
            # Validate file
            await self._validate_file(file_path)
            
            # Calculate file hash
            file_hash = self._calculate_file_hash(file_path)
            
            # Process based on file type
            suffix = file_path.suffix.lower()
            
            if suffix in [".pdf"]:
                content, tables, images = await self._process_pdf(file_path)
            elif suffix in [".docx", ".doc"]:
                content, tables, images = await self._process_word(file_path)
            elif suffix in [".xlsx", ".xls"]:
                content, tables, images = await self._process_excel(file_path)
            elif suffix in [".txt"]:
                content, tables, images = await self._process_text(file_path)
            elif suffix in [".jpg", ".png"]:
                content, tables, images = await self._process_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")
            
            # Combine metadata
            combined_metadata = {
                "filename": file_path.name,
                "file_type": suffix,
                "file_size": file_path.stat().st_size,
                "processing_time": time.time(),
                **(metadata or {})
            }
            
            return ProcessedDocument(
                content=content,
                metadata=combined_metadata,
                tables=tables,
                images=images,
                file_hash=file_hash
            )
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise

    async def _validate_file(self, file_path: Path) -> None:
        """Validate file before processing"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if file_path.suffix.lower() not in self.config.supported_formats:
            raise ValueError(f"Unsupported format: {file_path.suffix}")
            
        if file_path.stat().st_size > self.config.max_file_size_mb * 1024 * 1024:
            raise ValueError(f"File too large: {file_path}")

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        hasher = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hasher.update(chunk)
        return hasher.hexdigest()

    async def _process_pdf(
        self,
        file_path: Path
    ) -> tuple[str, List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process PDF document"""
        content_parts = []
        tables = []
        images = []
        
        doc = fitz.open(str(file_path))
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text
            content_parts.append(page.get_text())
            
            # Extract tables if enabled
            if self.config.extract_tables:
                page_tables = await self._extract_tables_from_pdf(page)
                tables.extend(page_tables)
            
            # Extract images
            page_images = await self._extract_images_from_pdf(page, page_num)
            images.extend(page_images)
            
            # Run OCR on images if enabled
            if self.config.ocr_enabled and page_images:
                ocr_text = await self._run_ocr_on_images(page_images)
                content_parts.append(ocr_text)
        
        doc.close()
        
        return "\n\n".join(content_parts), tables, images

    async def _process_word(
        self,
        file_path: Path
    ) -> tuple[str, List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process Word document"""
        doc = docx.Document(file_path)
        content_parts = []
        tables = []
        images = []
        
        # Extract text with formatting
        for paragraph in doc.paragraphs:
            if self.config.preserve_formatting:
                content_parts.append(self._format_paragraph(paragraph))
            else:
                content_parts.append(paragraph.text)
        
        # Extract tables
        if self.config.extract_tables:
            for table in doc.tables:
                table_data = self._extract_table_from_word(table)
                tables.append(table_data)
        
        # Extract images
        rel_ids = []
        for rel in doc.part.rels.values():
            if rel.reltype == docx.opc.constants.RELATIONSHIP_TYPE.IMAGE:
                image_data = {
                    "content": rel._target,
                    "type": rel.target_ref.split('.')[-1],
                    "source": "word_document"
                }
                images.append(image_data)
                
                if self.config.ocr_enabled:
                    ocr_text = await self._run_ocr_on_image(image_data["content"])
                    content_parts.append(ocr_text)
        
        return "\n\n".join(content_parts), tables, images

    async def _process_excel(
        self,
        file_path: Path
    ) -> tuple[str, List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process Excel document"""
        df_dict = pd.read_excel(file_path, sheet_name=None)
        content_parts = []
        tables = []
        
        for sheet_name, df in df_dict.items():
            # Extract text content
            content_parts.append(f"Sheet: {sheet_name}")
            content_parts.append(df.to_string())
            
            # Extract tables
            if self.config.extract_tables:
                table_data = {
                    "name": sheet_name,
                    "data": df.to_dict(orient="records"),
                    "columns": df.columns.tolist(),
                    "source": "excel_sheet"
                }
                tables.append(table_data)
        
        return "\n\n".join(content_parts), tables, []

    async def _process_text(
        self,
        file_path: Path
    ) -> tuple[str, List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process text document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content, [], []

    async def _process_image(
        self,
        file_path: Path
    ) -> tuple[str, List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process image file"""
        image = Image.open(file_path)
        image_data = {
            "content": image,
            "type": file_path.suffix[1:],
            "source": "image_file"
        }
        
        content = ""
        if self.config.ocr_enabled:
            content = pytesseract.image_to_string(image)
        
        return content, [], [image_data]

    async def _extract_tables_from_pdf(
        self,
        page: fitz.Page
    ) -> List[Dict[str, Any]]:
        """Extract tables from PDF page"""
        tables = []
        
        # Use external library like tabula-py or camelot
        # Placeholder for actual implementation
        return tables

    async def _extract_images_from_pdf(
        self,
        page: fitz.Page,
        page_num: int
    ) -> List[Dict[str, Any]]:
        """Extract images from PDF page"""
        images = []
        
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = page.parent.extract_image(xref)
            
            if base_image:
                image_data = {
                    "content": base_image["image"],
                    "type": base_image["ext"],
                    "source": f"pdf_page_{page_num}",
                    "index": img_index
                }
                images.append(image_data)
        
        return images

    async def _run_ocr_on_images(
        self,
        images: List[Dict[str, Any]]
    ) -> str:
        """Run OCR on list of images"""
        ocr_results = []
        
        for image_data in images:
            image_bytes = image_data["content"]
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image)
            if text.strip():
                ocr_results.append(text)
        
        return "\n\n".join(ocr_results)

    async def _run_ocr_on_image(self, image_bytes: bytes) -> str:
        """Run OCR on single image"""
        image = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(image)

    def _format_paragraph(self, paragraph: docx.text.paragraph.Paragraph) -> str:
        """Format Word paragraph preserving styles"""
        text = paragraph.text
        
        # Add basic formatting
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1])
            text = f"{'#' * level} {text}"
        
        if paragraph.style.font.bold:
            text = f"**{text}**"
        if paragraph.style.font.italic:
            text = f"*{text}*"
            
        return text

    def _extract_table_from_word(
        self,
        table: docx.table.Table
    ) -> Dict[str, Any]:
        """Extract data from Word table"""
        rows = []
        
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            rows.append(row_data)
            
        return {
            "data": rows,
            "num_rows": len(rows),
            "num_cols": len(rows[0]) if rows else 0,
            "source": "word_document"
        }

    def get_state(self) -> Dict[str, Any]:
        """Get current processor state"""
        return {
            "supported_formats": self.config.supported_formats,
            "ocr_enabled": self.config.ocr_enabled,
            "extract_tables": self.config.extract_tables,
            "max_file_size_mb": self.config.max_file_size_mb
        }